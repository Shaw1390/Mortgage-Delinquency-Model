import os
import json
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config import MODEL_DIR, REPORTS_DIR, TRAIN_YEARS, MONITOR_YEARS

NAVY = RGBColor(0x1F, 0x38, 0x64)
DARK = RGBColor(0x22, 0x22, 0x22)


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
    return h


def body(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK
    return p


def bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(10.5)
    return p


def add_metrics_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return table


if __name__ == "__main__":
    with open(os.path.join(MODEL_DIR, "metrics.json")) as f:
        metrics = json.load(f)

    drift_report = None
    governance_df = None
    drift_path = os.path.join(REPORTS_DIR, "drift_report.json")
    governance_path = os.path.join(REPORTS_DIR, "governance_status.csv")
    if os.path.exists(drift_path):
        with open(drift_path) as f:
            drift_report = json.load(f)
    if os.path.exists(governance_path):
        governance_df = pd.read_csv(governance_path)

    doc = Document()

    title = doc.add_heading("Model Card & Validation Report", level=0)
    for run in title.runs:
        run.font.color.rgb = NAVY
    subtitle = doc.add_paragraph("Mortgage Delinquency Risk Model \u2014 Logistic Regression")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.italic = True

    # 1. Model overview
    heading(doc, "1. Model Overview")
    body(doc, "Purpose: ", bold=True)
    body(doc, "Estimates the probability that a residential mortgage loan will reach 90+ "
              "days delinquent or REO status at any point over its life, using loan "
              "characteristics known at origination.")
    body(doc, "Intended use: ", bold=True)
    body(doc, "Decision-support for portfolio risk segmentation and monitoring. Intended "
              "to inform review prioritization and risk-tiering, not to serve as the sole "
              "basis for an individual credit decision.")
    body(doc, "Out of scope: ", bold=True)
    body(doc, "Not validated for use as an automated approve/decline system, and not "
              "calibrated for any population outside conforming conventional mortgages "
              "resembling the training data described below.")

    # 2. Data lineage
    heading(doc, "2. Data Lineage")
    bullet(doc, "Source: Freddie Mac Single-Family Loan-Level Dataset (public data), sample extracts")
    bullet(doc, f"Training vintages: {', '.join(str(y) for y in TRAIN_YEARS)} (pre-financial-crisis originations)")
    bullet(doc, f"Monitoring vintages: {', '.join(str(y) for y in MONITOR_YEARS)} (post-crisis, tighter underwriting)")
    bullet(doc, "Known limitation: this is U.S. loan-level data used as a stand-in for "
                "secured residential lending generally. Model risk management practices "
                "shown here follow general regulated-lending conventions rather than any "
                "specific institution's internal policy.")

    # 3. Validation methodology & metrics
    heading(doc, "3. Validation Methodology")
    body(doc, f"Trained on an 80/20 stratified split of the {', '.join(str(y) for y in TRAIN_YEARS)} "
              f"vintages ({metrics['n_train']:,} training loans, {metrics['n_validation']:,} held out "
              f"for validation). A HistGradientBoostingClassifier was trained as a benchmark; logistic "
              f"regression was selected as the documented model despite a lower AUC, prioritizing "
              f"coefficient-level interpretability for a regulated lending context.")

    lr = metrics["logistic_regression"]
    gbm = metrics["hist_gradient_boosting_benchmark"]
    add_metrics_table(
        doc,
        ["Metric", "Logistic Regression (selected)", "GBM (benchmark only)"],
        [
            ["AUC", lr["auc"], gbm["auc"]],
            ["Precision", lr["precision"], gbm["precision"]],
            ["Recall", lr["recall"], gbm["recall"]],
        ],
    )

    doc.add_paragraph()
    body(doc, f"Validation set default rate: {metrics['train_default_rate']}")

    # 4. Model risk tier
    heading(doc, "4. Model Risk Tier")
    body(doc, "Tier 2 \u2014 Decision-support model.", bold=True)
    body(doc, "Rationale: influences prioritization and monitoring workflows but does not "
              "directly execute credit decisions without human review. Would be reclassified "
              "to Tier 1 if adopted for automated decisioning.")

    # 5. Review cadence
    heading(doc, "5. Review Cadence")
    bullet(doc, "Scheduled revalidation: annually, or every 2 vintage years of new data, whichever is sooner")
    bullet(doc, "Off-cycle review triggers: any monitoring signal reaching \"Trigger review\" status "
                "(see Section 7), or a known shift in underwriting policy/economic conditions")

    # 6. Approval workflow
    heading(doc, "6. Approval Workflow")
    add_metrics_table(
        doc,
        ["Step", "Owner", "Action"],
        [
            ["1", "Model Developer", "Submits model, validation report, and monitoring plan"],
            ["2", "Model Risk / Governance Reviewer", "Independently reviews methodology and metrics"],
            ["3", "Business Owner", "Confirms intended use aligns with actual deployment"],
            ["4", "Governance Committee", "Approves for production monitoring use"],
        ],
    )

    # 7. Monitoring summary (only if phase 6 has run)
    heading(doc, "7. Monitoring Summary")
    if drift_report is not None and governance_df is not None:
        mp = drift_report["monitor_performance"]
        body(doc, f"Scored against the {', '.join(str(y) for y in MONITOR_YEARS)} monitor set "
                  f"(a regime the model never saw in training):")
        add_metrics_table(
            doc,
            ["Metric", "Validation", "Monitor Set", "Drop"],
            [["AUC", mp["validation_auc"], mp["auc"], mp["auc_drop"]]],
        )
        doc.add_paragraph()
        body(doc, "Governance status by monitoring signal:", bold=True)
        gov_table = add_metrics_table(
            doc,
            list(governance_df.columns),
            governance_df.values.tolist(),
        )
    else:
        body(doc, "Not yet available \u2014 run 06_monitoring_drift.py and re-generate this "
                  "report to populate this section with live drift and performance-decay results.")

    out_path = os.path.join(REPORTS_DIR, "Model_Card_Validation_Report.docx")
    os.makedirs(REPORTS_DIR, exist_ok=True)
    doc.save(out_path)
    print(f"Saved {out_path}")
